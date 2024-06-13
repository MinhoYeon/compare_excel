import pandas as pd
from docx import Document

def compare_excel_files(file1, file2):
    # 엑셀 파일 읽기
    df1 = pd.read_excel(file1)
    df2 = pd.read_excel(file2)
    
    # 동일한 식별 번호가 있는지 확인
    if not df1['순번'].equals(df2['순번']):
        raise ValueError("두 파일의 순번이 일치하지 않습니다.")
    
    # 각 필드의 값을 비교하여 동일한 셀의 데이터가 다른 셀 찾기
    differing_values_report = {}
    for column in df1.columns:
        if column == '순번':
            continue
        
        differing_values = df1[(df1[column] != df2[column]) & 
                               ~(df1[column].isna() & df2[column].isna())]
        
        if not differing_values.empty:
            differing_values_report[column] = differing_values[['순번', column]]
    
    return differing_values_report, df1, df2

def save_report_to_word(report, df1, df2, output_file):
    doc = Document()
    doc.add_heading('데이터 비교 보고서', level=1)
    
    for field, data in report.items():
        doc.add_heading(f'필드: {field}', level=2)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '순번'
        hdr_cells[1].text = '첫 번째 파일'
        hdr_cells[2].text = '두 번째 파일'
        
        for index, row in data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['순번'])
            row_cells[1].text = str(row[field])
            row_cells[2].text = str(df2.loc[index, field])
    
    doc.save(output_file)

# 파일 경로
file1 = './input/file1.xlsx'
file2 = './input/file2.xlsx'
output_file = './output/differing_values_report.docx'

# 비교 결과 저장
differing_values, df1, df2 = compare_excel_files(file1, file2)
save_report_to_word(differing_values, df1, df2, output_file)

print(f'비교 결과가 {output_file} 파일에 저장되었습니다.')

